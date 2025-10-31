# -*- coding: utf-8 -*-
"""
Gerador Automático de Cronogramas de Estudo (versão final com títulos hierárquicos e template .dotx)
- GUI (tkinter) para parâmetros e seleção de arquivos (inclui .dotx opcional).
- Leitura de lista_de_temas.xlsx e lista_de_aulas.xlsx (pandas, openpyxl).
- Filtro por tipo de prova com pesos; módulos com peso 0 são excluídos.
- Simulação diária com fases (Início/Meio/Final/Pré-prova) e cotas A/Q/R; empréstimos; carryover; resíduos.
- Revisão espaçada D+1,3,7,14,30 com realocação para o próximo dia de estudo.
- Remoção iterativa por prioridade se não couber (menor peso; empate maior carga horária).
- Geração de DOCX A4 (python-docx), com:
  * Capa (PNG).
  * Contracapa com rótulos em negrito e valores.
  * Orientações (aplica Heading 1 no título; incorpora DOCX/PNG; referencia PDF).
  * Cronograma semanal com hierarquia de títulos: Semana = Heading 1; Dia de estudo = Heading 2.
  * Checklist de módulos removidos quando abreviado.
- Template .dotx: documento pode ser criado a partir do template; no Windows, após salvar,
  os estilos do template são anexados e copiados via Word COM (se disponível).
- Conversão a PDF: tenta docx2pdf; se falhar, tenta Word COM; caso contrário mantém apenas DOCX.
- Memória de inputs em scheduler_config.json no diretório do script.

Dependências:
  pip install pandas openpyxl python-docx python-dateutil docx2pdf pywin32

Observações:
- Datas no formato DD/MM/AAAA.
- “Dias de estudo por semana” é numérico; dias fixos podem ser definidos nas configurações avançadas.
- Estrutura das planilhas conforme especificação.
"""
import subprocess  # para taskkill no Windows
from docx.shared import Pt, Cm, RGBColor
import os
import json
import math
import sys
import traceback
from datetime import datetime, timedelta, date
from collections import defaultdict, OrderedDict
from typing import Optional
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import FormulaRule

# Importações com mensagens claras caso faltem
try:
    import pandas as pd
except Exception as e:
    raise SystemExit("Instale pandas: pip install pandas") from e

try:
    from dateutil.relativedelta import relativedelta
except Exception as e:
    raise SystemExit("Instale python-dateutil: pip install python-dateutil") from e

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
except Exception as e:
    raise SystemExit("Tkinter é necessário (já vem no Python padrão em Windows).") from e

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENTATION
except Exception as e:
    raise SystemExit("Instale python-docx: pip install python-docx") from e

# docx2pdf é opcional
DOCX2PDF_AVAILABLE = False
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

WIN32_AVAILABLE = False
if os.name == "nt":
    try:
        import win32com.client  # type: ignore
        WIN32_AVAILABLE = True
    except Exception:
        WIN32_AVAILABLE = False

CONFIG_FILE = "scheduler_config.json"

TIPOS_PROVA = [
    "TEA","TSA","ME1","ME2","ME3",
    "ME1 1T","ME1 2T","ME1 3T","ME1 4T",
    "ME2 1T","ME2 2T","ME2 3T","ME2 4T",
    "ME3 1T","ME3 2T","ME3 3T","ME3 4T"
]

FRACTIONS_BY_PHASE = {
    "inicio":   {"A": 0.75, "Q": 0.15, "R": 0.10},  # primeiro quarto (EXCETO o Dia 1)
    "meio":     {"A": 0.65, "Q": 0.15, "R": 0.20},  # segundo quarto
    "final":    {"A": 0.60, "Q": 0.15, "R": 0.25},  # terceiro quarto
    "preprova": {"A": 0.50, "Q": 0.15, "R": 0.30}   # último quarto
}


#Essas constantes definem uma regra de "empréstimo" de tempo para acomodar aulas mais longas. O objetivo é evitar
#  que uma aula seja interrompida no meio só porque o tempo alocado para "Aulas" (A) acabou.
BORROW_Q_BY_PHASE = {
    "inicio": 0.40,
    "meio": 0.30,
    "final": 0.20,
    "preprova": 0.15
}
BORROW_R = 0.10


#Esta política define o que acontece com o tempo que "sobra" da cota de "Aulas" (A) em um dia. 
# Isso pode acontecer se, ao final do dia, não houver tempo suficiente para a próxima aula da fila.
RESIDUAL_POLICY = {
    "inicio":   {"Q": 1.00, "R": 0.00},
    "meio":     {"Q": 0.70, "R": 0.30},
    "final":    {"Q": 0.50, "R": 0.50},
    "preprova": {"Q": 0.30, "R": 0.70}
}

REVIEW_OFFSETS = [1, 3, 7, 14, 30]  # dias
DEFAULT_REVIEW_OFFSETS = [1, 3, 7, 14, 30, 90]  # valor padrão exibido na GUI

# --- FUNÇÕES AUXILIARES NOVAS ---

# --- Nomes de dias em português (0=Segunda ... 6=Domingo) ---
WEEKDAY_NAMES_PT = [
    "Segunda-feira", "Terça-feira", "Quarta-feira",
    "Quinta-feira", "Sexta-feira", "Sábado", "Domingo"
]


# ===== FUNÇÃO NOVA: exporta Excel no formato solicitado =====
def export_excel_schedule(out_xlsx_path, daily, study_days, exam_date):
    from datetime import timedelta
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.formatting.rule import FormulaRule

    def week_start(d):
        return d - timedelta(days=d.weekday())

    wb = Workbook()
    ws = wb.active
    ws.title = "Cronograma"

    ws["A1"] = "Data da prova"
    ws["B1"] = exam_date
    ws["C1"] = "Hoje"
    ws["D1"] = "=TODAY()"
    for c in ("A1", "B1", "C1", "D1"):
        ws[c].font = Font(bold=True)
        ws[c].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A3"

    headers = [
        "Semana","Tema","Aula","Data prevista","Aula assistida?","Questões respondidas?",
        "Desempenho","Revisão D1","Revisão D1 concluída?","Revisão D7","Data da Revisão D7",
        "Revisão D7 concluída?","Revisão D30","Data da Revisão D30","Revisão D30 concluída?",
        "Revisão D90","Data da Revisão D90","Revisão D90 concluída?"
    ]
    ws.append([""] * len(headers))
    for j, h in enumerate(headers, start=1):
        cell = ws.cell(row=2, column=j, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="404040")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    fill_green  = PatternFill(fill_type="solid", start_color="FFC6EFCE", end_color="FFC6EFCE")
    fill_red    = PatternFill(fill_type="solid", start_color="FFFFC7CE", end_color="FFFFC7CE")
    fill_yellow = PatternFill(fill_type="solid", start_color="FFFFF2CC", end_color="FFFFF2CC")
    fill_blue   = PatternFill(fill_type="solid", start_color="FFC6D9F1", end_color="FFC6D9F1")

    thin  = Side(style="thin",  color="D9D9D9")
    thick = Side(style="thick", color="000000")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)

    dv_desempenho = DataValidation(
        type="list",
        formula1='"(Inserir desempenho),≤ 60%,61-79%,≥ 80%"',
        allow_blank=True
    )
    ws.add_data_validation(dv_desempenho)

    def _week_index(d, d0):
        base = week_start(d0)
        return ((week_start(d) - base).days // 7) + 1

    def _set_border_edges(c, left=None, right=None, top=None, bottom=None):
        c.border = Border(
            left=left     if left   is not None else c.border.left,
            right=right   if right  is not None else c.border.right,
            top=top       if top    is not None else c.border.top,
            bottom=bottom if bottom is not None else c.border.bottom
        )

    rows_start = 3
    row = rows_start
    week_to_first_last = {}

    first_day = study_days[0] if study_days else None
    for d in study_days:
        lessons = daily[d]["A_lessons"]
        if not lessons:
            continue
        wnum = _week_index(d, first_day)
        for lesson in lessons:
            tema = lesson["modulo"]
            aula = lesson["aula"]

            ws.cell(row=row, column=1, value=f"Semana {wnum}")
            ws.cell(row=row, column=2, value=tema)
            ws.cell(row=row, column=3, value=aula)

            c_data = ws.cell(row=row, column=4, value=d)
            c_data.number_format = "dd/mm/yyyy"

            ce = ws.cell(row=row, column=5, value=False); ce.number_format = ";;;"
            cf = ws.cell(row=row, column=6, value=False); cf.number_format = ";;;"

            cg = ws.cell(row=row, column=7, value="(Inserir desempenho)")
            dv_desempenho.add(cg)

            ws.cell(
                row=row, column=8,
                value='=IF(LOWER($G{r})="(inserir desempenho)","",IF($G{r}="","",IF(LEFT($G{r},2)="≤ ","Recomendada","Não recomendada")))'.format(r=row)
            )
            ci = ws.cell(row=row, column=9, value=False); ci.number_format = ";;;"

            ws.cell(
                row=row, column=10,
                value='=IF(LOWER($G{r})="(inserir desempenho)","",IF($G{r}="","",IF(OR($G{r}="≤ 60%",$G{r}="61-79%"),"Recomendada","Não recomendada")))'.format(r=row)
            )
            # K: Data D7 com regras de valor incluindo "Não recomendada" se J="Não recomendada"
            ck = ws.cell(
                row=row, column=11,
                value='=IF(LOWER($G{r})="(inserir desempenho)","",IF($J{r}="Não recomendada","Não recomendada",IF($G{r}="≥ 80%","",IF($D{r}="","",$D{r}+7))))'.format(r=row)
            )
            ck.number_format = "dd/mm/yyyy"
            cl = ws.cell(row=row, column=12, value=""); cl.number_format = ";;;"

            ws.cell(row=row, column=13, value="Recomendada")
            cn = ws.cell(row=row, column=14, value='=IF($D{r}="","",$D{r}+30)'.format(r=row))
            cn.number_format = "dd/mm/yyyy"
            co = ws.cell(row=row, column=15, value=""); co.number_format = ";;;"

            ws.cell(row=row, column=16, value='=IF($D{r}+90 < $B$1,"Recomendada","Não recomendada")'.format(r=row))
            cq = ws.cell(row=row, column=17, value='=IF($D{r}="","",$D{r}+90)'.format(r=row))
            cq.number_format = "dd/mm/yyyy"
            cr = ws.cell(row=row, column=18, value=""); cr.number_format = ";;;"

            for col in range(1, 19):
                c = ws.cell(row=row, column=col)
                c.border = border_all
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

            if wnum not in week_to_first_last:
                week_to_first_last[wnum] = [row, row]
            else:
                week_to_first_last[wnum][1] = row

            row += 1

    last_row = row - 1

    for wnum, (r1, r2) in week_to_first_last.items():
        if r2 > r1:
            ws.merge_cells(start_row=r1, start_column=1, end_row=r2, end_column=1)
        ws.cell(row=r1, column=1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.cell(row=r1, column=1).font = Font(bold=True)

    for wnum, (r1, r2) in week_to_first_last.items():
        r = r1
        while r <= r2:
            tema_atual = ws.cell(row=r, column=2).value
            r_end = r
            while r_end + 1 <= r2 and ws.cell(row=r_end + 1, column=2).value == tema_atual:
                r_end += 1
            if r_end > r:
                ws.merge_cells(start_row=r, start_column=2, end_row=r_end, end_column=2)
                bcell = ws.cell(row=r, column=2)
                bcell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            r = r_end + 1

    for wnum, (r1, r2) in week_to_first_last.items():
        for col in range(1, 19):
            _set_border_edges(ws.cell(row=r1, column=col), top=thick)
            _set_border_edges(ws.cell(row=r2, column=col), bottom=thick)
        for r in range(r1, r2 + 1):
            _set_border_edges(ws.cell(row=r, column=1), left=thick)
            _set_border_edges(ws.cell(row=r, column=18), right=thick)

    if last_row >= 3:
        for r in range(3, last_row + 1):
            ws.cell(row=r, column=1).fill = fill_yellow
            ws.cell(row=r, column=2).fill = fill_yellow
            ws.cell(row=r, column=3).fill = fill_yellow
            ws.cell(row=r, column=4).fill = fill_yellow

        rng_e = f"E3:E{last_row}"
        marcado_expr_e = 'OR($E3=TRUE,$E3=1,$E3="TRUE",$E3="VERDADEIRO")'
        ws.conditional_formatting.add(rng_e, FormulaRule(formula=[marcado_expr_e], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_e, FormulaRule(formula=[f'AND(NOT({marcado_expr_e}),$D3<>"",$D3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_e, FormulaRule(formula=[f'AND(NOT({marcado_expr_e}),OR($D3>=$D$1,$D3=""))'], fill=fill_yellow))

        rng_f = f"F3:F{last_row}"
        marcado_expr_f = 'OR($F3=TRUE,$F3=1,$F3="TRUE",$F3="VERDADEIRO")'
        ws.conditional_formatting.add(rng_f, FormulaRule(formula=[marcado_expr_f], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_f, FormulaRule(formula=[f'AND(NOT({marcado_expr_f}),$D3<>"",$D3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_f, FormulaRule(formula=[f'AND(NOT({marcado_expr_f}),OR($D3>=$D$1,$D3=""))'], fill=fill_yellow))

        ws.conditional_formatting.add(f"G3:G{last_row}", FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))
        ws.conditional_formatting.add(f"G3:G{last_row}", FormulaRule(formula=['$G3="≤ 60%"'], fill=fill_red))
        ws.conditional_formatting.add(f"G3:G{last_row}", FormulaRule(formula=['$G3="61-79%"'], fill=fill_yellow))
        ws.conditional_formatting.add(f"G3:G{last_row}", FormulaRule(formula=['$G3="≥ 80%"'], fill=fill_green))

        rng_h = f"H3:H{last_row}"
        ws.conditional_formatting.add(rng_h, FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))
        ws.conditional_formatting.add(rng_h, FormulaRule(formula=['AND($H3="Recomendada", NOT(OR($I3=TRUE,$I3=1,$I3="TRUE",$I3="VERDADEIRO")))'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_h, FormulaRule(formula=['AND($H3="Recomendada", OR($I3=TRUE,$I3=1,$I3="TRUE",$I3="VERDADEIRO"))'], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_h, FormulaRule(formula=['$H3="Não recomendada"'], fill=fill_green, stopIfTrue=True))

        ws.conditional_formatting.add(f"I3:I{last_row}", FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))

        rng_k = f"K3:K{last_row}"
        marcado_expr_l = 'OR($L3=TRUE,$L3=1,$L3="TRUE",$L3="VERDADEIRO")'
        ws.conditional_formatting.add(rng_k, FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))
        ws.conditional_formatting.add(rng_k, FormulaRule(formula=['$J3="Não recomendada"'], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_k, FormulaRule(formula=['$G3="≥ 80%"'], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_k, FormulaRule(formula=[f'AND(OR($G3="61-79%",$G3="≤ 60%"),NOT({marcado_expr_l}))'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_k, FormulaRule(formula=[f'AND(OR($G3="61-79%",$G3="≤ 60%"),{marcado_expr_l})'], fill=fill_green, stopIfTrue=True))

        # Concluída: manter apenas H/I e J/L
        def add_done_cfmt(col_recom, col_done):
            ws.conditional_formatting.add(f"{col_done}3:{col_done}{last_row}", FormulaRule(formula=[f'AND(${col_recom}3="Recomendada",${col_done}3<>TRUE)'], fill=fill_red))
            ws.conditional_formatting.add(f"{col_done}3:{col_done}{last_row}", FormulaRule(formula=[f'OR(AND(${col_recom}3="Recomendada",${col_done}3=TRUE),${col_recom}3="Não recomendada")'], fill=fill_green))
        add_done_cfmt("H","I")
        add_done_cfmt("J","L")

        ws.conditional_formatting.add(f"J3:J{last_row}", FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))
        ws.conditional_formatting.add(f"J3:J{last_row}", FormulaRule(formula=['$J3="Recomendada"'], fill=fill_red))
        ws.conditional_formatting.add(f"J3:J{last_row}", FormulaRule(formula=['$J3="Não recomendada"'], fill=fill_green))

        ws.conditional_formatting.add(f"L3:L{last_row}", FormulaRule(formula=['LOWER($G3)="(inserir desempenho)"'], fill=fill_yellow, stopIfTrue=True))

        # M: sempre recomendada; coloração por N (prazo) e O (checklist)
        rng_m = f"M3:M{last_row}"
        marcado_expr_o = 'OR($O3=TRUE,$O3=1,$O3="TRUE",$O3="VERDADEIRO")'
        ws.conditional_formatting.add(rng_m, FormulaRule(formula=[marcado_expr_o], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_m, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),$N3<>"",$N3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_m, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),OR($N3>=$D$1,$N3=""))'], fill=fill_yellow, stopIfTrue=True))

        # N: coloração por O e prazo N
        rng_n = f"N3:N{last_row}"
        ws.conditional_formatting.add(rng_n, FormulaRule(formula=[marcado_expr_o], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_n, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),$N3<>"",$N3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_n, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),OR($N3>=$D$1,$N3=""))'], fill=fill_yellow, stopIfTrue=True))

        # O: coloração por O e prazo N
        rng_o = f"O3:O{last_row}"
        ws.conditional_formatting.add(rng_o, FormulaRule(formula=[marcado_expr_o], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_o, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),$N3<>"",$N3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_o, FormulaRule(formula=[f'AND(NOT({marcado_expr_o}),OR($N3>=$D$1,$N3=""))'], fill=fill_yellow, stopIfTrue=True))

        # P e R: coloração por R e prazo Q
        rng_p = f"P3:P{last_row}"
        rng_q = f"Q3:Q{last_row}"
        rng_r = f"R3:R{last_row}"
        marcado_expr_r = 'OR($R3=TRUE,$R3=1,$R3="TRUE",$R3="VERDADEIRO")'
        ws.conditional_formatting.add(rng_p, FormulaRule(formula=[marcado_expr_r], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_p, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),$Q3<>"",$Q3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_p, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),OR($Q3>=$D$1,$Q3=""))'], fill=fill_yellow, stopIfTrue=True))

        ws.conditional_formatting.add(rng_r, FormulaRule(formula=[marcado_expr_r], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_r, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),$Q3<>"",$Q3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_r, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),OR($Q3>=$D$1,$Q3=""))'], fill=fill_yellow, stopIfTrue=True))

        ws.conditional_formatting.add(rng_q, FormulaRule(formula=[marcado_expr_r], fill=fill_green, stopIfTrue=True))
        ws.conditional_formatting.add(rng_q, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),$Q3<>"",$Q3<$D$1)'], fill=fill_red, stopIfTrue=True))
        ws.conditional_formatting.add(rng_q, FormulaRule(formula=[f'AND(NOT({marcado_expr_r}),OR($Q3>=$D$1,$Q3=""))'], fill=fill_yellow, stopIfTrue=True))

    for r in (1, 2):
        for c in range(1, 19):
            ws.cell(row=r, column=c).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws.cell(row=r, column=c).border = border_all

    for col in range(1, 19):
        _set_border_edges(ws.cell(row=2, column=col), top=thick, bottom=thick)
    _set_border_edges(ws.cell(row=2, column=1), left=thick)
    _set_border_edges(ws.cell(row=2, column=18), right=thick)

    ws.column_dimensions["A"].width = 11
    ws.column_dimensions["B"].width = 34
    ws.column_dimensions["C"].width = 41
    ws.column_dimensions["D"].width = 11
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 14
    ws.column_dimensions["G"].width = 16
    ws.column_dimensions["H"].width = 15
    ws.column_dimensions["I"].width = 12
    ws.column_dimensions["J"].width = 13
    ws.column_dimensions["K"].width = 13  # solicitado
    ws.column_dimensions["L"].width = 12
    ws.column_dimensions["M"].width = 13  # solicitado
    ws.column_dimensions["N"].width = 11
    ws.column_dimensions["O"].width = 12
    ws.column_dimensions["P"].width = 13  # solicitado
    ws.column_dimensions["Q"].width = 11
    ws.column_dimensions["R"].width = 12

    wb.save(out_xlsx_path)

    try:
        import os as _os
        import win32com.client  # type: ignore
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        wbcom = excel.Workbooks.Open(_os.path.abspath(out_xlsx_path))
        sh = wbcom.Worksheets("Cronograma")

        sh.Rows(f"3:{last_row}").EntireRow.AutoFit()

        def _clear_checkboxes_in_range(col_letter):
            try:
                for cb in list(sh.CheckBoxes()):
                    tl = cb.TopLeftCell
                    if tl is not None and tl.Column == sh.Range(col_letter + "1").Column:
                        if 3 <= tl.Row <= last_row:
                            cb.Delete()
            except Exception:
                pass

        def _add_checkboxes_centered(col_letter):
            col_ref = sh.Range(f"{col_letter}1")
            col_left  = float(col_ref.Left)
            col_width = float(col_ref.Width)
            col_center = col_left + col_width / 2.0
            sh.Range(f"{col_letter}3:{col_letter}{last_row}").NumberFormat = ";;;"

            heights = []
            for r in range(3, last_row + 1):
                rh = sh.Rows(r).RowHeight
                if rh is None or rh == 0:
                    rh = sh.Rows(r).Height
                heights.append(float(rh))
            heights.sort()
            median_h = heights[len(heights)//2] if heights else float(sh.Rows(3).Height)

            base = min(col_width, median_h)
            cb_size = max(base * 0.68, 11.0)
            cb_size = round(cb_size, 0)

            for r in range(3, last_row + 1):
                rng = sh.Range(f"{col_letter}{r}")
                rh = sh.Rows(r).RowHeight
                if rh is None or rh == 0:
                    rh = sh.Rows(r).Height
                left = round(col_center - cb_size / 2.0, 0)
                top  = round(float(rng.Top) + (float(rh) - cb_size) / 2.0, 0)
                cb = sh.CheckBoxes().Add(left, top, cb_size, cb_size)
                cb.Caption = ""
                cb.LinkedCell = f"{col_letter}{r}"
                cb.Value = 0
                try:
                    cb.AutoSize = False
                    cb.Placement = 2
                    cb.PrintObject = True
                except Exception:
                    pass

        if last_row >= 3:
            for col_letter in ("E", "F", "I", "L", "O", "R"):
                _clear_checkboxes_in_range(col_letter)
                _add_checkboxes_centered(col_letter)

        wbcom.Close(SaveChanges=True)
        excel.Quit()
    except Exception:
        from openpyxl import load_workbook as _load
        wb2 = _load(out_xlsx_path)
        ws2 = wb2["Cronograma"]
        ws2.column_dimensions["A"].width = 11
        ws2.column_dimensions["B"].width = 34
        ws2.column_dimensions["C"].width = 41
        ws2.column_dimensions["D"].width = 11
        ws2.column_dimensions["E"].width = 10
        ws2.column_dimensions["F"].width = 14
        ws2.column_dimensions["G"].width = 16
        ws2.column_dimensions["H"].width = 15
        ws2.column_dimensions["I"].width = 12
        ws2.column_dimensions["J"].width = 13
        ws2.column_dimensions["K"].width = 13
        ws2.column_dimensions["L"].width = 12
        ws2.column_dimensions["M"].width = 13
        ws2.column_dimensions["N"].width = 11
        ws2.column_dimensions["O"].width = 12
        ws2.column_dimensions["P"].width = 13
        ws2.column_dimensions["Q"].width = 11
        ws2.column_dimensions["R"].width = 12
        wb2.save(out_xlsx_path)

    return out_xlsx_path

# ===== FIM DA FUNÇÃO NOVA =====

def format_day_with_name(d: date) -> str:
    # Retorna "Segunda-feira (DD/MM)"
    return f"{WEEKDAY_NAMES_PT[d.weekday()]} ({d.strftime('%d/%m')})"

def kill_office_processes():
    """
    Encerra processos residuais do Excel e do Word ao final da execução.
    Usa 'taskkill' no Windows; silencioso e tolerante a erros.
    Observação: isso FINALIZA QUALQUER instância aberta (inclusive as abertas manualmente).
    """
    if os.name != "nt":
        return
    for proc in ("EXCEL.EXE", "WINWORD.EXE"):
        try:
            subprocess.run(
                ["taskkill", "/F", "/IM", proc, "/T"],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
        except Exception:
            pass

def _resolve_orient_source(orient_path: str | None) -> Optional[str]:
    """
    Resolve a fonte das orientações segundo a regra:
      1) Se o caminho fornecido existir, usa-o.
      2) Se não existir, tenta 'revisao_espacada_orientacoes.docx' no diretório do script.
      3) Se ainda não encontrar, abre GUI para o usuário selecionar PDF, DOCX ou PNG,
         e memoriza o caminho escolhido em scheduler_config.json.
    """
    try_path = orient_path if orient_path else ""
    if try_path and os.path.isfile(try_path):
        return try_path

    here = os.path.abspath(os.path.dirname(__file__))
    cand_docx = os.path.join(here, "revisao_espacada_orientacoes.docx")
    if os.path.isfile(cand_docx):
        return cand_docx

    # Abre GUI de seleção e memoriza o último input do usuário
    try:
        # cria root oculto para filedialog
        root = tk.Tk()
        root.withdraw()
        chosen = filedialog.askopenfilename(
            title="Selecione as orientações (PDF, DOCX ou PNG)",
            filetypes=[("PDF", "*.pdf"), ("Word", "*.docx"), ("Imagem (PNG)", "*.png")],
        )
        root.destroy()
        if chosen and os.path.isfile(chosen):
            cfg = load_config()
            cfg["orient_path"] = chosen
            save_config(cfg)
            return chosen
    except Exception:
        pass

    return None


# --- SUBSTITUA A FUNÇÃO POR ESTA VERSÃO COM FULL-BLEED ---

def _insert_pdf_as_images(doc: Document, pdf_path: str, full_bleed: bool = False):
    """
    Insere todas as páginas do PDF como imagens.
    Em full_bleed, cada página é colocada em sua própria seção sem margens,
    e ao final as margens originais são restauradas em uma nova seção.
    """
    pages_png = []
    page_sizes_in = []
    def _begin_full_bleed_section(doc: Document, w_in: float, h_in: float):
        base = doc.sections[-1]
        snapshot = {
            "page_width": base.page_width,
            "page_height": base.page_height,
            "top_margin": base.top_margin,
            "bottom_margin": base.bottom_margin,
            "left_margin": base.left_margin,
            "right_margin": base.right_margin,
            "header_distance": base.header_distance,
            "footer_distance": base.footer_distance,
        }
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width = Inches(w_in)
        sec.page_height = Inches(h_in)
        zero = Inches(0)
        sec.top_margin = zero
        sec.bottom_margin = zero
        sec.left_margin = zero
        sec.right_margin = zero
        sec.header_distance = zero
        sec.footer_distance = zero
        return snapshot

    def _end_full_bleed_section(doc: Document, snapshot: dict):
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width = snapshot["page_width"]
        sec.page_height = snapshot["page_height"]
        sec.top_margin = snapshot["top_margin"]
        sec.bottom_margin = snapshot["bottom_margin"]
        sec.left_margin = snapshot["left_margin"]
        sec.right_margin = snapshot["right_margin"]
        sec.header_distance = snapshot["header_distance"]
        sec.footer_distance = snapshot["footer_distance"]

    try:
        import fitz  # PyMuPDF
        with fitz.open(pdf_path) as pdf:
            for page in pdf:
                w_in = float(page.rect.width) / 72.0
                h_in = float(page.rect.height) / 72.0
                page_sizes_in.append((w_in, h_in))
                dpi = 216
                zoom = dpi / 72.0
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                tmp.write(pix.tobytes("png"))
                tmp.flush()
                tmp.close()
                pages_png.append(tmp.name)
    except Exception:
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path, dpi=216)
            a4_w_in, a4_h_in = 8.27, 11.69
            for img in images:
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                img.save(tmp.name, format="PNG")
                pages_png.append(tmp.name)
                page_sizes_in.append((a4_w_in, a4_h_in))
        except Exception:
            pages_png = []

    if not pages_png:
        doc.add_paragraph(f"Não foi possível incorporar o PDF. Consulte o arquivo em: {pdf_path}")
        return

    restore_snapshot = None

    for idx, png_path in enumerate(pages_png):
        w_in, h_in = page_sizes_in[idx] if idx < len(page_sizes_in) else (8.27, 11.69)

        if full_bleed:
            if restore_snapshot is None:
                restore_snapshot = {
                    "page_width": doc.sections[-1].page_width,
                    "page_height": doc.sections[-1].page_height,
                    "top_margin": doc.sections[-1].top_margin,
                    "bottom_margin": doc.sections[-1].bottom_margin,
                    "left_margin": doc.sections[-1].left_margin,
                    "right_margin": doc.sections[-1].right_margin,
                    "header_distance": doc.sections[-1].header_distance,
                    "footer_distance": doc.sections[-1].footer_distance,
                }
            _ = _begin_full_bleed_section(doc, w_in, h_in)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png_path, width=Inches(w_in))
        except Exception:
            doc.add_paragraph(f"[Falha ao inserir a imagem renderizada da página {idx+1} do PDF]")

        if not full_bleed and idx < len(pages_png) - 1:
            doc.add_page_break()

    if full_bleed and restore_snapshot is not None:
        _end_full_bleed_section(doc, restore_snapshot)

    for png_path in pages_png:
        try:
            os.remove(png_path)
        except Exception:
            pass

def _insert_docx_preserving_basic_layout(doc: Document, src_docx_path: str):
    """
    Copia o conteúdo de um DOCX preservando diagramação básica: headings, parágrafos, alinhamento,
    e propriedades de fonte dos runs (negrito, itálico, sublinhado, tamanho e nome de fonte).
    Observação: elementos muito avançados (seções, colunas, quebras complexas, imagens internas)
    podem não ser reproduzidos integralmente via python-docx.
    """
    from docx import Document as DocxReader
    src = DocxReader(src_docx_path)

    for p_src in src.paragraphs:
        p_dst = doc.add_paragraph()
        try:
            if p_src.style and p_src.style.name in [s.name for s in doc.styles]:
                p_dst.style = doc.styles[p_src.style.name]
        except Exception:
            pass
        try:
            p_dst.alignment = p_src.alignment
        except Exception:
            pass
        for r_src in p_src.runs:
            r_dst = p_dst.add_run(r_src.text)
            try:
                r_dst.bold = r_src.bold
                r_dst.italic = r_src.italic
                r_dst.underline = r_src.underline
                if r_src.font and r_src.font.size:
                    r_dst.font.size = r_src.font.size
                if r_src.font and r_src.font.name:
                    r_dst.font.name = r_src.font.name
            except Exception:
                pass

    for t_src in src.tables:
        rows = len(t_src.rows)
        cols = len(t_src.columns)
        t_dst = doc.add_table(rows=rows, cols=cols)
        try:
            t_dst.style = t_src.style
        except Exception:
            pass
        for i in range(rows):
            for j in range(cols):
                cell_text = "\n".join(p.text for p in t_src.cell(i, j).paragraphs)
                t_dst.cell(i, j).text = cell_text
        doc.add_paragraph("")

def add_orientacoes(doc: Document, orient_path: Optional[str]):
    # Incorpora as orientações diretamente, sem título prévio.
    # Suporta PDF (full-bleed), DOCX (parágrafos/tabelas) e PNG (imagem centrada).
    try:
        resolved = _resolve_orient_source(orient_path)
        if not resolved or not os.path.isfile(resolved):
            doc.add_paragraph("Arquivo de orientações não encontrado. Prossiga consultando o material externo.")
            doc.add_page_break()
            return

        ext = os.path.splitext(resolved)[1].lower()
        if ext == ".pdf":
            _insert_pdf_as_images(doc, resolved, full_bleed=True)
        elif ext == ".docx":
            _insert_docx_preserving_basic_layout(doc, resolved)
        elif ext == ".png":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            try:
                r.add_picture(resolved, width=Cm(16))
            except Exception:
                doc.add_paragraph(f"[Falha ao inserir PNG de orientações: {resolved}]")
        else:
            doc.add_paragraph(f"Formato de orientações não suportado: {resolved}")

        doc.add_page_break()
    except Exception as e:
        try:
            doc.add_paragraph(f"[Falha ao incorporar orientações: {e}]")
            doc.add_page_break()
        except Exception:
            pass


def set_page_background(doc: Document, hex_color: str = "000000"):
    doc_elm = doc._element
    existing = doc_elm.find(qn("w:background"))
    if existing is not None:
        doc_elm.remove(existing)
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    doc_elm.insert(0, bg)

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_config(cfg: dict):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def parse_date_br(s: str) -> date:
    return datetime.strptime(s.strip(), "%d/%m/%Y").date()

def format_date_br(d: date) -> str:
    return d.strftime("%d/%m/%Y")

def week_start(d: date) -> date:
    # Domingo como início de semana
    return d - timedelta(days=(d.weekday() + 1) % 7)

def generate_study_days(dini: date, dprova: date, dias_por_semana: int, custom_weekdays=None):
    # custom_weekdays: conjunto de inteiros 0..6 para seg..dom; se None, usa primeiros N dias no intervalo semanal
    if dias_por_semana <= 0:
        return []
    study_days = []
    cur = dini
    while cur <= dprova:
        start = week_start(cur)
        week_days = []
        if custom_weekdays:
            for i in range(7):
                candidate = start + timedelta(days=i)
                if candidate < dini or candidate > dprova:
                    continue
                if candidate.weekday() in custom_weekdays:  # 0=Seg … 6=Dom
                    week_days.append(candidate)
        else:
            for i in range(7):
                candidate = start + timedelta(days=i)
                if candidate < dini or candidate > dprova:
                    continue
                week_days.append(candidate)
            week_days.sort()
            week_days = week_days[:dias_por_semana]
        study_days.extend([d for d in week_days if d >= dini and d <= dprova])
        cur = start + timedelta(days=7)
    study_days = sorted(list(OrderedDict.fromkeys(study_days)))
    return study_days

def determine_phase(idx: int, total_days: int) -> str:
    if total_days == 0:
        return "inicio"
    q = total_days // 4
    if idx < q:
        return "inicio"
    elif idx < 2*q:
        return "meio"
    elif idx < 3*q:
        return "final"
    else:
        return "preprova"

def normalize_reviews(reviews_raw, study_days):
    from collections import defaultdict
    study_days_set = set(study_days)
    last = max(study_days)
    reviews = defaultdict(list)

    for t_raw, items in reviews_raw.items():
        # Casos com data de revisão DEPOIS do último dia de estudo: omite do DOCX/PDF.
        if t_raw > last:
            continue

        # Data de revisão exatamente em um dia de estudo: mantém.
        if t_raw in study_days_set:
            reviews[t_raw].extend(items)
            continue

        # Senão, envia para o PRÓXIMO dia de estudo disponível em ou após t_raw, sem ultrapassar 'last'.
        d = t_raw
        while d <= last and d not in study_days_set:
            d = d + timedelta(days=1)
        if d in study_days_set:
            reviews[d].extend(items)
        # Se não existir um dia de estudo até 'last', a revisão é omitida.

    return reviews

def next_study_day_on_or_after(target: date, study_days_set):
    d = target
    while True:
        if d in study_days_set:
            return d
        d = d + timedelta(days=1)

def load_document_with_template(template_path: Optional[str]) -> Document:
    if template_path and os.path.isfile(template_path):
        try:
            return Document(template_path)
        except Exception:
            return Document()
    return Document()

def apply_template_styles_win(docx_path: str, template_path: str) -> bool:
    if not (WIN32_AVAILABLE and os.path.isfile(template_path) and os.path.isfile(docx_path)):
        return False
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.AttachedTemplate = os.path.abspath(template_path)
        doc.UpdateStylesOnOpen = True
        doc.Close(SaveChanges=True)
        wdOrganizerObjectStyles = 3
        src = os.path.abspath(template_path)
        dst = os.path.abspath(docx_path)
        core_styles = [
            "Title","Heading 1","Heading 2","Heading 3","Heading 4","Heading 5","Heading 6","Heading 7","Heading 8","Heading 9",
            "Título","Título 1","Título 2","Título 3","Título 4","Título 5","Título 6","Título 7","Título 8","Título 9"
        ]
        for s in core_styles:
            try:
                word.OrganizerCopy(Source=src, Destination=dst, Name=s, Object=wdOrganizerObjectStyles)
            except Exception:
                pass
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.Save()
        doc.Close()
        word.Quit()
        return True
    except Exception:
        try:
            word.Quit()
        except Exception:
            pass
        return False

def read_inputs_from_gui(prefill: dict):
    root = tk.Tk()
    root.title("Gerador de Cronograma – Parâmetros")

    minutos_var = tk.StringVar(value=str(prefill.get("minutos_por_dia","180")))
    dias_semana_var = tk.StringVar(value=str(prefill.get("dias_por_semana","5")))
    data_inicio_var = tk.StringVar(value=prefill.get("data_inicio","01/11/2025"))
    data_prova_var = tk.StringVar(value=prefill.get("data_prova","31/12/2025"))
    tipo_var = tk.StringVar(value=prefill.get("tipo_prova","TEA"))

    temas_path_var = tk.StringVar(value=prefill.get("temas_path",""))
    aulas_path_var = tk.StringVar(value=prefill.get("aulas_path",""))
    capa_path_var = tk.StringVar(value=prefill.get("capa_path",""))
    orient_path_var = tk.StringVar(value=prefill.get("orient_path",""))
    template_path_var = tk.StringVar(value=prefill.get("template_path",""))

    advanced_open = tk.BooleanVar(value=False)
    weekday_vars = [tk.BooleanVar(value=False) for _ in range(7)]  # seg..dom

    # NOVO: seleção de offsets de revisão
    preselected_offsets = prefill.get("review_offsets", DEFAULT_REVIEW_OFFSETS)
    review_vars = {d: tk.BooleanVar(value=(d in preselected_offsets)) for d in DEFAULT_REVIEW_OFFSETS}

    def browse_excel(var):
        path = filedialog.askopenfilename(title="Selecione o arquivo Excel", filetypes=[("Excel","*.xlsx")])
        if path:
            var.set(path)

    def browse_png(var):
        path = filedialog.askopenfilename(title="Selecione a capa (PNG)", filetypes=[("PNG","*.png")])
        if path:
            var.set(path)

    def browse_orient(var):
        path = filedialog.askopenfilename(title="Selecione orientações (PDF, DOCX ou PNG)", filetypes=[("Arquivos","*.pdf *.docx *.png")])
        if path:
            var.set(path)

    def browse_dotx(var):
        path = filedialog.askopenfilename(title="Selecione o arquivo de estilos (.dotx)", filetypes=[("Modelo do Word","*.dotx")])
        if path:
            var.set(path)

    frm = ttk.Frame(root, padding=10)
    frm.grid(row=0, column=0, sticky="nsew")

    ttk.Label(frm, text="Minutos de estudo por dia").grid(row=0, column=0, sticky="w")
    ttk.Entry(frm, textvariable=minutos_var, width=12).grid(row=0, column=1, sticky="w")

    ttk.Label(frm, text="Dias de estudo por semana").grid(row=1, column=0, sticky="w")
    ttk.Entry(frm, textvariable=dias_semana_var, width=12).grid(row=1, column=1, sticky="w")

    ttk.Label(frm, text="Data de início (DD/MM/AAAA)").grid(row=2, column=0, sticky="w")
    ttk.Entry(frm, textvariable=data_inicio_var, width=12).grid(row=2, column=1, sticky="w")

    ttk.Label(frm, text="Data da prova (DD/MM/AAAA)").grid(row=3, column=0, sticky="w")
    ttk.Entry(frm, textvariable=data_prova_var, width=12).grid(row=3, column=1, sticky="w")

    ttk.Label(frm, text="Tipo de prova").grid(row=4, column=0, sticky="w")
    ttk.Combobox(frm, textvariable=tipo_var, values=TIPOS_PROVA, width=12, state="readonly").grid(row=4, column=1, sticky="w")

    ttk.Separator(frm).grid(row=5, column=0, columnspan=3, sticky="ew", pady=(8,8))

    ttk.Label(frm, text="lista_de_temas.xlsx").grid(row=6, column=0, sticky="w")
    ttk.Entry(frm, textvariable=temas_path_var, width=48).grid(row=6, column=1, sticky="w")
    ttk.Button(frm, text="Procurar", command=lambda: browse_excel(temas_path_var)).grid(row=6, column=2, sticky="w")

    ttk.Label(frm, text="lista_de_aulas.xlsx").grid(row=7, column=0, sticky="w")
    ttk.Entry(frm, textvariable=aulas_path_var, width=48).grid(row=7, column=1, sticky="w")
    ttk.Button(frm, text="Procurar", command=lambda: browse_excel(aulas_path_var)).grid(row=7, column=2, sticky="w")

    ttk.Label(frm, text="Capa (PNG)").grid(row=8, column=0, sticky="w")
    ttk.Entry(frm, textvariable=capa_path_var, width=48).grid(row=8, column=1, sticky="w")
    ttk.Button(frm, text="Procurar", command=lambda: browse_png(capa_path_var)).grid(row=8, column=2, sticky="w")

    ttk.Label(frm, text="Orientações (PDF/DOCX/PNG)").grid(row=9, column=0, sticky="w")
    ttk.Entry(frm, textvariable=orient_path_var, width=48).grid(row=9, column=1, sticky="w")
    ttk.Button(frm, text="Procurar", command=lambda: browse_orient(orient_path_var)).grid(row=9, column=2, sticky="w")

    ttk.Label(frm, text="Arquivo de estilos (.dotx)").grid(row=10, column=0, sticky="w")
    ttk.Entry(frm, textvariable=template_path_var, width=48).grid(row=10, column=1, sticky="w")
    ttk.Button(frm, text="Procurar", command=lambda: browse_dotx(template_path_var)).grid(row=10, column=2, sticky="w")

    def toggle_advanced():
        val = not advanced_open.get()
        advanced_open.set(val)
        if val:
            adv_frame.grid()
        else:
            adv_frame.grid_remove()

    adv_button = ttk.Button(frm, text="Configurações avançadas (dias fixos da semana)", command=toggle_advanced)
    adv_button.grid(row=11, column=0, columnspan=3, sticky="w", pady=(8,4))

    adv_frame = ttk.Frame(frm, padding=(10,6))
    adv_frame.grid(row=12, column=0, columnspan=3, sticky="w")
    adv_frame.grid_remove()

    weekdays_labels = ["Seg","Ter","Qua","Qui","Sex","Sáb","Dom"]
    for i, lab in enumerate(weekdays_labels):
        ttk.Checkbutton(adv_frame, text=lab, variable=weekday_vars[i]).grid(row=0, column=i, sticky="w")

    # NOVO BLOCO: seleção dos intervalos de revisão
    review_frame = ttk.Frame(frm, padding=(0,6))
    review_frame.grid(row=13, column=0, columnspan=3, sticky="w")
    ttk.Label(review_frame, text="Intervalos de revisão (dias):").grid(row=0, column=0, sticky="w", pady=(6,2))

    # exibição horizontal: 1, 3, 7, 14, 30
    for col, d in enumerate(DEFAULT_REVIEW_OFFSETS, start=0):
        ttk.Checkbutton(review_frame, text=str(d), variable=review_vars[d]).grid(row=1, column=col, sticky="w")

    def on_ok():
        try:
            minutos = int(minutos_var.get())
            dias_semana = int(dias_semana_var.get())
            di = parse_date_br(data_inicio_var.get())
            dp = parse_date_br(data_prova_var.get())
            if minutos <= 0 or dias_semana <= 0:
                messagebox.showerror("Erro", "Minutos e dias/semana devem ser positivos.")
                return
            if dp < di:
                messagebox.showerror("Erro", "Data da prova não pode ser anterior à data de início.")
                return
            if not os.path.isfile(temas_path_var.get()):
                messagebox.showerror("Erro", "Selecione lista_de_temas.xlsx.")
                return
            if not os.path.isfile(aulas_path_var.get()):
                messagebox.showerror("Erro", "Selecione lista_de_aulas.xlsx.")
                return
            if not os.path.isfile(capa_path_var.get()):
                messagebox.showerror("Erro", "Selecione a capa PNG.")
                return

            selected_offsets = sorted([d for d, v in review_vars.items() if v.get()])

            root.result = {
                "minutos_por_dia": minutos,
                "dias_por_semana": dias_semana,
                "data_inicio": di,
                "data_prova": dp,
                "tipo_prova": tipo_var.get(),
                "temas_path": temas_path_var.get(),
                "aulas_path": aulas_path_var.get(),
                "capa_path": capa_path_var.get(),
                "orient_path": orient_path_var.get(),
                "template_path": template_path_var.get().strip(),
                "custom_weekdays": {i for i, v in enumerate(weekday_vars) if v.get()},
                "review_offsets": selected_offsets
            }
            root.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Entrada inválida: {e}")

    ttk.Button(frm, text="Gerar", command=on_ok).grid(row=14, column=0, pady=(12,0))
    ttk.Button(frm, text="Cancelar", command=root.destroy).grid(row=14, column=1, pady=(12,0))

    root.mainloop()
    if hasattr(root, "result"):
        return root.result
    else:
        raise SystemExit("Cancelado pelo usuário.")

def read_dataframes(temas_path, aulas_path):
    temas = pd.read_excel(temas_path, engine="openpyxl")
    aulas = pd.read_excel(aulas_path, engine="openpyxl")

    temas_required = ["Nome do Tema"] + TIPOS_PROVA
    for col in temas_required:
        if col not in temas.columns:
            raise ValueError(f"lista_de_temas.xlsx não contém a coluna obrigatória: {col}")

    aulas_required = ["Nome da Aula","Nome do Tema","Duração"]
    for col in aulas_required:
        if col not in aulas.columns:
            raise ValueError(f"lista_de_aulas.xlsx não contém a coluna obrigatória: {col}")

    aulas["Duração"] = aulas["Duração"].astype(int)
    return temas, aulas

def build_lessons_queue(temas_df, aulas_df, tipo_prova):
    # 1) Calcula pesos por tipo de prova e elimina módulos com peso 0
    temas_df = temas_df.copy()
    temas_df["peso"] = temas_df[tipo_prova].astype(int)
    temas_valid = temas_df[temas_df["peso"] > 0].copy()

    # 2) Custo total por módulo (só para relatórios e desempate na remoção)
    custos = (
        aulas_df.groupby("Nome do Tema")["Duração"]
        .sum()
        .reset_index()
        .rename(columns={"Duração": "custo"})
    )
    temas_valid = temas_valid.merge(custos, on="Nome do Tema", how="left").fillna({"custo": 0})

    # 3) Mapas de apoio
    peso_map = dict(zip(temas_valid["Nome do Tema"], temas_valid["peso"]))
    custo_map = dict(zip(temas_valid["Nome do Tema"], temas_valid["custo"]))

    # 4) Fila de aulas NA ORDEM DO EXCEL, limitada aos módulos com peso > 0
    mod_valid_set = set(temas_valid["Nome do Tema"])
    aulas_filtradas = aulas_df[aulas_df["Nome do Tema"].isin(mod_valid_set)]

    lessons = []
    for _, row in aulas_filtradas.iterrows():
        modulo = str(row["Nome do Tema"])
        lessons.append({
            "aula": str(row["Nome da Aula"]),
            "modulo": modulo,
            "dur": int(row["Duração"]),
            "peso": int(peso_map.get(modulo, 0)),
        })

    # 5) Apenas para exibição/relatório: módulos hierarquizados por peso (não afeta a alocação)
    mod_order = (
        temas_valid.sort_values(by=["peso", "custo"], ascending=[False, True])["Nome do Tema"].tolist()
    )

    return lessons, peso_map, custo_map, mod_order
# --- AJUSTAR a assinatura do simulador e usar offsets escolhidos pelo usuário ---
def simulate_schedule(study_days, minutos_dia, lessons_all, peso_map, review_offsets):
    from collections import defaultdict, OrderedDict
    study_days_set = set(study_days)
    total_days = len(study_days)

    daily = OrderedDict()
    for d in study_days:
        daily[d] = {"A_lessons": [], "Q_min": 0, "R_min": 0,
                    "phase": determine_phase(study_days.index(d), total_days)}

    reviews_raw = defaultdict(list)
    queue = list(lessons_all)

    must_force_carryover = False

    for idx, d in enumerate(study_days):
        phase = daily[d]["phase"]

        if idx == 0:
            fr = {"A": 0.80, "Q": 0.20, "R": 0.00}
        else:
            fr = FRACTIONS_BY_PHASE[phase]

        A_quota = minutos_dia * fr["A"]
        Q_quota = minutos_dia * fr["Q"]
        R_quota = minutos_dia * fr["R"]

        max_borrow_Q = Q_quota * BORROW_Q_BY_PHASE[phase]
        max_borrow_R = R_quota * BORROW_R

        borrowed_Q = 0.0
        borrowed_R = 0.0
        force_debt = 0.0

        def _force_first_if_needed():
            nonlocal A_quota, Q_quota, R_quota, borrowed_Q, borrowed_R, force_debt, must_force_carryover
            if not queue:
                must_force_carryover = False
                return
            lesson = queue.pop(0)
            dur = float(lesson["dur"])

            use_A = min(A_quota, dur)
            A_quota -= use_A
            remain = dur - use_A

            avail_Qb = max(0.0, max_borrow_Q - borrowed_Q)
            use_Qb = min(avail_Qb, remain)
            borrowed_Q += use_Qb
            remain -= use_Qb

            avail_Rb = max(0.0, max_borrow_R - borrowed_R)
            use_Rb = min(avail_Rb, remain)
            borrowed_R += use_Rb
            remain -= use_Rb

            if remain > 1e-6:
                force_debt += remain

            daily[d]["A_lessons"].append(lesson)

            for off in review_offsets:
                t_raw = d + timedelta(days=off)
                reviews_raw[t_raw].append({
                    "aula": lesson["aula"],
                    "modulo": lesson["modulo"],
                    "watched_date": d,
                    "peso": int(peso_map.get(lesson["modulo"], 0))
                })

            must_force_carryover = False

        if must_force_carryover:
            _force_first_if_needed()

        while queue:
            dur = float(queue[0]["dur"])
            available = A_quota + max(0.0, max_borrow_Q - borrowed_Q) + max(0.0, max_borrow_R - borrowed_R)
            if dur <= available + 1e-6:
                need = max(0.0, dur - A_quota)

                take_Q = min(need, max(0.0, max_borrow_Q - borrowed_Q))
                borrowed_Q += take_Q
                need -= take_Q

                take_R = min(need, max(0.0, max_borrow_R - borrowed_R))
                borrowed_R += take_R
                need -= take_R

                A_quota -= max(0.0, dur - (take_Q + take_R))
                if A_quota < 0.0:
                    A_quota = 0.0

                lesson = queue.pop(0)
                daily[d]["A_lessons"].append(lesson)

                for off in review_offsets:
                    t_raw = d + timedelta(days=off)
                    reviews_raw[t_raw].append({
                        "aula": lesson["aula"],
                        "modulo": lesson["modulo"],
                        "watched_date": d,
                        "peso": int(peso_map.get(lesson["modulo"], 0))
                    })
            else:
                must_force_carryover = True
                break

        if not daily[d]["A_lessons"] and queue:
            _force_first_if_needed()

        resid = A_quota
        pol = RESIDUAL_POLICY[phase]
        Q_quota += resid * pol["Q"]
        R_quota += resid * pol["R"]
        A_quota = 0.0

        Q_final = max(0, int(round(Q_quota - borrowed_Q - force_debt / 2.0)))
        R_final = max(0, int(round(R_quota - borrowed_R - force_debt / 2.0)))

        daily[d]["Q_min"] = Q_final
        daily[d]["R_min"] = R_final

    all_allocated = (len(queue) == 0)
    reviews = normalize_reviews(reviews_raw, study_days)
    return all_allocated, daily, reviews, queue


def try_fit_with_removals(study_days, minutos_dia, lessons_all, peso_map, custo_map, review_offsets):
    ok, daily, reviews, remaining = simulate_schedule(study_days, minutos_dia, lessons_all, peso_map, review_offsets)
    if ok:
        return True, daily, reviews, []

    mod_info = {}
    for lesson in lessons_all:
        m = lesson["modulo"]
        mod_info.setdefault(m, {"peso": peso_map.get(m,0), "custo":0})
        mod_info[m]["custo"] += lesson["dur"]

    mods_sorted = sorted(mod_info.items(), key=lambda kv: (kv[1]["peso"], -kv[1]["custo"]))
    removed_modules = []

    working_lessons = list(lessons_all)
    for m, meta in mods_sorted:
        working_lessons = [l for l in working_lessons if l["modulo"] != m]
        removed_modules.append(m)
        ok, daily, reviews, remaining = simulate_schedule(study_days, minutos_dia, working_lessons, peso_map, review_offsets)
        if ok:
            removed_lessons = [l for l in lessons_all if l["modulo"] in removed_modules]
            return True, daily, reviews, removed_lessons

    removed_lessons = [l for l in lessons_all if l["modulo"] in removed_modules]
    return False, daily, reviews, removed_lessons

def ensure_a4(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

def add_cover(doc, capa_path: str):
    # 1) Zera margens da PRIMEIRA seção para permitir full-bleed
    sec0 = doc.sections[0]
    orig = {
        "page_width": sec0.page_width,
        "page_height": sec0.page_height,
        "top": sec0.top_margin,
        "bottom": sec0.bottom_margin,
        "left": sec0.left_margin,
        "right": sec0.right_margin,
        "h": sec0.header_distance,
        "f": sec0.footer_distance,
    }
    zero = Cm(0)
    sec0.top_margin = zero
    sec0.bottom_margin = zero
    sec0.left_margin = zero
    sec0.right_margin = zero
    sec0.header_distance = zero
    sec0.footer_distance = zero

    # 2) Insere a CAPA ocupando 100% da página (largura e altura da página)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(
        capa_path,
        width=orig["page_width"],
        height=orig["page_height"]
    )

    # 3) Abre nova seção e RESTAURA margens normais para o restante do documento
    sec_rest = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_rest.page_width  = orig["page_width"]
    sec_rest.page_height = orig["page_height"]
    sec_rest.top_margin    = Cm(2.0)
    sec_rest.bottom_margin = Cm(2.0)
    sec_rest.left_margin   = Cm(2.0)
    sec_rest.right_margin  = Cm(2.0)
    sec_rest.header_distance = Cm(1.25)
    sec_rest.footer_distance = Cm(1.25)

def add_contracapa(doc: Document, tipo_prova, di, dp, min_dia, dps, total_weeks, total_A_min, total_QR_min, completo: bool, removed_count):
    h = doc.add_paragraph(f"Tipo de prova: {tipo_prova}")
    try:
        h.style = doc.styles["Heading 1"]
    except Exception:
        pass
    
    # LINHA ADICIONADA: solta um parágrafo normal logo após o título
    p_blank = doc.add_paragraph()
    try:
        p_blank.style = doc.styles["Normal"]
    except Exception:
        pass

    # Parágrafo Normal em branco
    p_after_title = doc.add_paragraph()
    try:
        p_after_title.style = doc.styles["Normal"]
    except Exception:
        pass

    h = doc.add_paragraph(f"Especificações Personalizadas")
    try:
        h.style = doc.styles["Heading 3"]
    except Exception:
        pass

    # LINHA ADICIONADA: solta um parágrafo normal logo após o título
    p_blank = doc.add_paragraph()
    try:
        p_blank.style = doc.styles["Normal"]
    except Exception:
        pass

    def add_label_value(label, value, sufixo=""):
        p = doc.add_paragraph()
        r_label = p.add_run(f"{label}: ")
        r_label.bold = True
        p.add_run(f"{value}{sufixo}")

    add_label_value("Data de início", format_date_br(di))
    add_label_value("Data da prova", format_date_br(dp))
    add_label_value("Minutos de estudo por dia", min_dia)
    add_label_value("Dias de estudo por semana", dps)
    add_label_value("Duração total do cronograma em semanas", total_weeks)
    add_label_value("Tempo total de aulas programadas", total_A_min, " minutos")
    add_label_value("Tempo total de questões + revisão", total_QR_min, " minutos")

    # Parágrafo Normal em branco
    p_after_title = doc.add_paragraph()
    try:
        p_after_title.style = doc.styles["Normal"]
    except Exception:
        pass

    # Parágrafo Normal em branco
    p_after_title = doc.add_paragraph()
    try:
        p_after_title.style = doc.styles["Normal"]
    except Exception:
        pass

    # TÍTULO INSERIDO ENTRE AS LINHAS SOLICITADAS
    titulo_secao = doc.add_paragraph("Tipo de Cronograma")
    try:
        titulo_secao.style = doc.styles["Heading 3"]
    except Exception:
        pass
    # Parágrafo Normal em branco para espaçamento após o título
    p_after_title = doc.add_paragraph()
    try:
        p_after_title.style = doc.styles["Normal"]
    except Exception:
        pass


    if completo:
        doc.add_paragraph("Cronograma Completo.")
    else:
        doc.add_paragraph("Cronograma Abreviado.")
        p = doc.add_paragraph()
        r = p.add_run("Nota: ")
        r.bold = True
        p.add_run(f"{removed_count} aulas foram removidas por limitação de capacidade. A lista detalhada consta ao final do documento; é facultado ao aluno realizar substituições manuais conforme domínio individual dos temas.")

    doc.add_page_break()

def iter_weeks(study_days):
    if not study_days:
        return
    current_week = week_start(study_days[0])
    buf = []
    for d in study_days:
        ws = week_start(d)
        if ws != current_week and buf:
            yield current_week, buf
            buf = []
            current_week = ws
        buf.append(d)
    if buf:
        yield current_week, buf

# AJUSTE: adicionar parâmetro label_dates para controlar exibição de datas nos dias de estudo
def add_schedule(doc: Document, study_days, daily, reviews, peso_map, label_dates: bool):
    for wstart, days in iter_weeks(study_days):
        # AJUSTE: semana sempre com 7 dias (segunda a domingo)
        wend = wstart + timedelta(days=6)
        p_week = doc.add_paragraph("Semana {} a {}".format(format_date_br(wstart), format_date_br(wend)))
        try:
            p_week.style = doc.styles["Heading 1"]
        except Exception:
            pass

        dia_count = 0
        for d in days:
            dia_count += 1
            node = daily[d]
            # AJUSTE: se não houver dias fixos selecionados, não exibir data específica
            if label_dates:
                # Exibe: "Dia 1 - Segunda-feira (DD/MM)"
                p_day = doc.add_paragraph(f"Dia {dia_count} - {format_day_with_name(d)}")
            else:
                p_day = doc.add_paragraph(f"Dia de estudo {dia_count}")
            try:
                p_day.style = doc.styles["Heading 2"]
            except Exception:
                pass

            if not label_dates:
                doc.add_paragraph("Cronograma finalizado. Você pode alocar esse tempo para assistir aulas recém lançadas na plataforma ou expandir sua revisão.")

            aulas = node["A_lessons"]
            total_aulas_min = sum(l["dur"] for l in aulas)

            #Aulas para Assistir
            h_aulas = doc.add_paragraph("Aulas para Assistir   ({} min)".format(total_aulas_min))
            try:
                h_aulas.style = doc.styles["Heading 3"]
            except Exception:
                pass
            for l in aulas:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)   # remove espaço entre parágrafos
                p.paragraph_format.space_before = Pt(0)  # idem antes
                run_marker = p.add_run("➙ ")
                run_marker.font.color.rgb = RGBColor(217, 187, 38)
                run_marker.bold = True
                p.add_run("{} - {} min".format(l["aula"], l["dur"]))

            #Treinamento de Questões
            h_q = doc.add_paragraph("Treinamento de Questões   ({} min)".format(node["Q_min"]))
            try:
                h_q.style = doc.styles["Heading 3"]
            except Exception:
                pass
            doc.add_paragraph("Resolução de exercícios referentes às aulas do dia.")

            #REVISÃO ESPAÇADA
            h_r = doc.add_paragraph("Revisão Espaçada   ({} min)".format(node["R_min"]))
            try:
                h_r.style = doc.styles["Heading 3"]
            except Exception:
                pass
            todays_reviews = sorted(reviews.get(d, []), key=lambda x: (-x["peso"], x["modulo"], x["aula"]))
            if todays_reviews:
                for item in todays_reviews:
                    # calcula "há X dias" com singular/plural e caso "hoje"
                    days_ago = (d - item["watched_date"]).days
                    if days_ago <= 0:
                        quando = "hoje"
                    elif days_ago == 1:
                        quando = "há 1 dia"
                    else:
                        quando = f"há {days_ago} dias"
                    doc.add_paragraph(f"{item['aula']} (Assistida {quando}).")
            else:
                doc.add_paragraph("Sem itens de revisão programados para hoje.")
            doc.add_paragraph("")


        doc.add_page_break()

def add_removed_checklist(doc: Document, removed_lessons):
    if not removed_lessons:
        return
    h = doc.add_paragraph("Checklist de módulos removidos")
    try:
        h.style = doc.styles["Heading 1"]
    except Exception:
        pass
    by_mod = defaultdict(list)
    for l in removed_lessons:
        by_mod[l["modulo"]].append(l)
    for m in sorted(by_mod.keys()):
        doc.add_paragraph(m)
        lessons = by_mod[m]
        for l in lessons:
            doc.add_paragraph(" - {} ({} min)".format(l["aula"], l["dur"]))

def compute_totals(daily):
    total_A = 0
    total_QR = 0
    for d, node in daily.items():
        total_A += sum(l["dur"] for l in node["A_lessons"])
        total_QR += node["Q_min"] + node["R_min"]
    return total_A, total_QR

def export_to_pdf(docx_path: str):
    import os, time

    docx_abs = os.path.abspath(docx_path)
    pdf_abs  = os.path.splitext(docx_abs)[0] + ".pdf"

    def _win_long(p: str) -> str:
        # Word/COM tem limitações com caminhos longos; prefixo \\?\ ajuda em > 260 chars
        return p if os.name != "nt" or len(p) < 240 else "\\\\?\\" + p

    # 1) Windows/COM: caminho feliz e mais estável
    if WIN32_AVAILABLE:
        try:
            import win32com.client  # type: ignore
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                # Suprime prompts do Word (ex.: confirmação de conversão)
                word.DisplayAlerts = 0
            except Exception:
                pass

            doc = word.Documents.Open(_win_long(docx_abs))
            try:
                # SaveAs2 é mais robusto para PDF do que ExportAsFixedFormat em alguns ambientes
                # 17 = wdFormatPDF
                doc.SaveAs2(_win_long(pdf_abs), FileFormat=17)
            finally:
                # Fecha doc independentemente de sucesso
                doc.Close(SaveChanges=False)
                # Encerra a instância criada por nós
                word.Quit()

            # Espera curta e síncrona até o PDF existir e ter tamanho > 0
            for _ in range(60):
                if os.path.exists(pdf_abs) and os.path.getsize(pdf_abs) > 0:
                    return pdf_abs
                time.sleep(0.1)
            return None
        except Exception:
            # Se COM falhar por qualquer motivo, cai para docx2pdf se disponível
            pass

    # 2) Fallback docx2pdf: passe o diretório de saída, não o arquivo
    if DOCX2PDF_AVAILABLE:
        try:
            out_dir = os.path.dirname(pdf_abs) or "."
            # Alguns builds exigem diretório; ele cria <nome>.pdf automaticamente
            docx2pdf_convert(docx_abs, out_dir)
            candidate = os.path.join(out_dir, os.path.basename(pdf_abs))
            return candidate if os.path.exists(candidate) else None
        except Exception:
            return None

    # 3) Sem COM e sem docx2pdf
    return None

# --- PERSISTIR a escolha dos offsets no config em main() e repassar adiante ---
def main():
    cfg = load_config()
    params = read_inputs_from_gui(cfg)

    cfg.update({
        "minutos_por_dia": params["minutos_por_dia"],
        "dias_por_semana": params["dias_por_semana"],
        "data_inicio": format_date_br(params["data_inicio"]),
        "data_prova": format_date_br(params["data_prova"]),
        "tipo_prova": params["tipo_prova"],
        "temas_path": params["temas_path"],
        "aulas_path": params["aulas_path"],
        "capa_path": params["capa_path"],
        "orient_path": params["orient_path"],
        "template_path": params.get("template_path",""),
        "custom_weekdays": sorted(list(params["custom_weekdays"])) if params["custom_weekdays"] else [],
        "review_offsets": params.get("review_offsets", DEFAULT_REVIEW_OFFSETS)
    })
    save_config(cfg)

    if not params.get("template_path"):
        here = os.path.abspath(os.path.dirname(__file__))
        default_tpl = os.path.join(here, "Estilo.dotx")
        if os.path.isfile(default_tpl):
            params["template_path"] = default_tpl

    temas_df, aulas_df = read_dataframes(params["temas_path"], params["aulas_path"])
    lessons_all, peso_map, custo_map, mod_order = build_lessons_queue(temas_df, aulas_df, params["tipo_prova"])

    custom_weekdays = set(cfg.get("custom_weekdays", []))
    study_days = generate_study_days(params["data_inicio"], params["data_prova"], params["dias_por_semana"], custom_weekdays if custom_weekdays else None)
    if not study_days:
        raise SystemExit("Não há dias de estudo dentro do intervalo fornecido.")

    # NOVO: extrair offsets selecionados (pode estar vazio)
    review_offsets = params.get("review_offsets", DEFAULT_REVIEW_OFFSETS)

    ok, daily, reviews, removed_lessons = try_fit_with_removals(
        study_days, params["minutos_por_dia"], lessons_all, peso_map, custo_map, review_offsets
    )
    completo = ok and len(removed_lessons) == 0

    total_A_min, total_QR_min = compute_totals(daily)

    first_ws = week_start(study_days[0])
    last_ws = week_start(study_days[-1])
    total_weeks = ((last_ws - first_ws).days // 7) + 1

    out_base = "Cronograma_{}_{}_{}_{}xS_{}min".format(
        "Completo" if completo else "Abreviado",
        params["tipo_prova"].replace(" ",""),
        params["data_inicio"].strftime("%Y-%m-%d") + "_" + params["data_prova"].strftime("%Y-%m-%d"),
        params["dias_por_semana"],
        params["minutos_por_dia"]
    ).replace(":", "-")

    out_docx = out_base + ".docx"

    doc = load_document_with_template(params.get("template_path"))
    set_page_background(doc, "000000")
    ensure_a4(doc)

    add_cover(doc, params["capa_path"])
    add_contracapa(doc, params["tipo_prova"], params["data_inicio"], params["data_prova"],
                   params["minutos_por_dia"], params["dias_por_semana"], total_weeks,
                   total_A_min, total_QR_min, completo, len(removed_lessons))

    label_dates = bool(custom_weekdays)
    add_orientacoes(doc, params["orient_path"])
    add_schedule(doc, study_days, daily, reviews, peso_map, label_dates)
    if not completo:
        add_removed_checklist(doc, removed_lessons)

    doc.save(out_docx)

    tpl = params.get("template_path")
    if tpl:
        apply_template_styles_win(out_docx, tpl)

    pdf_path = export_to_pdf(out_docx)

    # === SOMENTE exportar Excel se o PDF foi confirmado ===
    out_xlsx = None
    if pdf_path:
        out_xlsx = out_base + ".xlsx"
        export_excel_schedule(out_xlsx, daily, study_days, params["data_prova"])

    msg = ["Cronograma gerado com sucesso."]
    msg.append("Arquivo DOCX: {}".format(os.path.abspath(out_docx)))
    if pdf_path:
        msg.append("Arquivo PDF: {}".format(os.path.abspath(pdf_path)))
    else:
        msg.append("PDF não gerado automaticamente. Instale docx2pdf ou utilize Microsoft Word no Windows para converter.")
    if out_xlsx:
        msg.append("Arquivo XLSX: {}".format(os.path.abspath(out_xlsx)))
    try:
        messagebox.showinfo("Concluído", "\n".join(msg))
    except Exception:
        pass


if __name__ == "__main__":
    try:
        main()
    except SystemExit as se:
        print(str(se))
    except Exception as e:
        traceback.print_exc()
        try:
            messagebox.showerror("Erro fatal", str(e))
        except Exception:
            pass
    finally:
        try:
            kill_office_processes()
        except Exception:
            pass